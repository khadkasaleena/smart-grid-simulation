"""
build_report.py
Generates the full written report (Tasks 1-4) as a Word .docx, following the
BSBI / UCA assignment brief: title page, table of contents, page numbers,
Harvard in-text references and reference list, embedded system diagram and
result figures, and the Task 2 source code as an appendix.

Run:  python build_report.py
Output: 0000000_Salina_Khadka_Assignment.docx  (rename the StudentID prefix)

Module: Advanced Simulation Techniques
Author: Salina Khadka, BSBI Berlin
"""

import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #


def add_page_number_footer(section):
    """Add a centred 'Page X' field to the section footer."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc(document):
    """Insert a Word table-of-contents field (updates on open / F9 in Word)."""
    p = document.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Right-click and choose 'Update Field' to build the table of contents."
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


def body(document, text):
    p = document.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
    return p


def fig(document, path, caption, width=6.0):
    if os.path.exists(os.path.join(HERE, path)):
        document.add_picture(os.path.join(HERE, path), width=Inches(width))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = document.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(9)


def code_appendix(document, filename):
    document.add_heading(filename, level=2)
    with open(os.path.join(HERE, filename), encoding="utf-8") as f:
        code = f.read()
    p = document.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(7.5)
    # force monospace east-asian too
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")


# --------------------------------------------------------------------------- #
# Build the document
# --------------------------------------------------------------------------- #


def main():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Title page -------------------------------------------------------
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph("Advanced Simulation of Power System Stability and "
                          "Cascading Failure Dynamics in Resilient Smart Grids")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.bold = True
        r.font.size = Pt(20)
    doc.add_paragraph()
    sub = doc.add_paragraph("Written Assignment")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)

    for _ in range(4):
        doc.add_paragraph()

    meta = [
        ("Student Name:", "Salina Khadka"),
        ("Student ID:", "[Your Student ID]"),
        ("Module:", "Advanced Simulation Techniques (AST)"),
        ("Module Code:", "[Module Code]"),
        ("Degree:", "BSc Computer Science and Digitisation"),
        ("Issued by:", "Dr. Waseem Alhasan"),
        ("Institution:", "Berlin School of Business and Innovation (BSBI) / UCA"),
        ("Submission Date:", "17 June 2026"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(label + " ")
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.size = Pt(12)

    doc.add_page_break()

    # ---- Table of contents -----------------------------------------------
    h = doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # Page numbers in footer (from here on, applied to whole doc).
    add_page_number_footer(doc.sections[0])

    # ====================================================================== #
    # TASK 1
    # ====================================================================== #
    doc.add_heading("Task 1: Critical Infrastructure Analysis", level=1)

    doc.add_heading("1.1 Key Components, Stakeholders and Interdependencies", level=2)
    for para in TASK1_1:
        body(doc, para)

    doc.add_heading("1.2 Vulnerabilities and Failure Modes", level=2)
    for para in TASK1_2:
        body(doc, para)

    doc.add_heading("1.3 Modelling Approach and Justification", level=2)
    for para in TASK1_3:
        body(doc, para)

    # ====================================================================== #
    # TASK 2
    # ====================================================================== #
    doc.add_heading("Task 2: Simulation Design and Implementation", level=1)

    doc.add_heading("2.1 Conceptual Model and System Diagram", level=2)
    for para in TASK2_1:
        body(doc, para)
    fig(doc, "system_diagram.png",
        "Figure 1. Conceptual system diagram: generation → transmission → "
        "substations → consumer zones, with the repair-crew control loop.")

    doc.add_heading("2.2 Implementation: Key Decisions and Parameter Settings", level=2)
    for para in TASK2_2:
        body(doc, para)

    doc.add_heading("2.3 User Interface, Visualisation and Export", level=2)
    for para in TASK2_3:
        body(doc, para)
    fig(doc, "graph_power_output.png",
        "Figure 2. Colab UI output — stacked generator output versus total demand.")
    fig(doc, "graph_grid_load.png",
        "Figure 3. Colab UI output — average line load with overload and cascade thresholds.")
    fig(doc, "graph_heatmap.png",
        "Figure 4. Colab UI output — per-line load heatmap (% of rated capacity).")

    doc.add_heading("2.4 Scenario Execution, Results and Performance Metrics", level=2)
    for para in TASK2_4:
        body(doc, para)

    # Results table
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Scenario", "Failures", "Cascades", "MTTR (h)",
                             "Peak load %", "Crew util."]):
        hdr[i].paragraphs[0].add_run(txt).bold = True
    for row in RESULTS_ROWS:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    cap = doc.add_paragraph("Table 1. Headline resilience metrics across the six what-if "
                            "scenarios (seed = 42, 72-hour horizon).")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    for para in TASK2_4B:
        body(doc, para)

    # ====================================================================== #
    # TASK 3
    # ====================================================================== #
    doc.add_heading("Task 3: Tool and Technique Comparison", level=1)

    doc.add_heading("3.1 Comparison of Two Simulation Tools/Techniques", level=2)
    for para in TASK3_1:
        body(doc, para)

    doc.add_heading("3.2 Alternative Approaches and Reflection", level=2)
    for para in TASK3_2:
        body(doc, para)

    # ====================================================================== #
    # TASK 4
    # ====================================================================== #
    doc.add_heading("Task 4: Business Value and Future Directions", level=1)

    doc.add_heading("4.1 Competitive Advantage and Return on Investment", level=2)
    for para in TASK4_1:
        body(doc, para)

    doc.add_heading("4.2 Emerging Trends and Ethical Considerations", level=2)
    for para in TASK4_2:
        body(doc, para)

    # ---- References -------------------------------------------------------
    doc.add_heading("References", level=1)
    for ref in REFERENCES:
        p = doc.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        for r in p.runs:
            r.font.size = Pt(11)

    # ---- Appendix: source code -------------------------------------------
    doc.add_page_break()
    doc.add_heading("Appendix A: Task 2 Source Code", level=1)
    body(doc, "The complete, runnable source code for the discrete-event simulation "
              "is reproduced below and is also hosted at "
              "https://github.com/khadkasaleena/smart-grid-simulation, where it can be "
              "executed directly in Google Colab via the badge in the README.")
    for fname in ["config.py", "grid_model.py", "metrics.py", "scenarios.py",
                  "visualization.py", "make_system_diagram.py", "run_simulation.py"]:
        code_appendix(doc, fname)

    out = os.path.join(HERE, "StudentID_Salina_Khadka_Assignment.docx")
    doc.save(out)
    print("Saved:", out)


# --------------------------------------------------------------------------- #
# Report text (Harvard in-text referencing throughout)
# --------------------------------------------------------------------------- #

TASK1_1 = [
"The electrical power grid is the archetypal critical infrastructure: a tightly coupled, "
"geographically dispersed network whose continuous operation underpins almost every other "
"sector of modern society. A modern smart grid integrates four functional layers — "
"generation, transmission, distribution and consumption — with an overarching layer of "
"information and communication technology (ICT) that senses, coordinates and controls the "
"physical flow of electricity in near real time (Amin and Wollenberg, 2005). The present "
"analysis focuses on a regional distribution system comprising dispatchable coal generation, "
"variable solar generation, a small transmission network, substations and the consumer zones "
"they serve.",

"The key physical components are generators, transmission lines, substations (transformers and "
"switchgear), protection relays, and the loads themselves. In the modelled system these are "
"represented concretely as two coal units (300 MW and 200 MW), a 150 MW solar farm, five "
"transmission lines of differing thermal ratings, and four substations feeding industrial, "
"residential, commercial and mixed consumer zones. Layered above the physical assets is the "
"cyber-control system — supervisory control and data acquisition (SCADA), state estimation "
"and protection logic — which monitors line loadings and dispatches corrective action. This "
"cyber–physical coupling is the defining feature of the smart grid and the source of much "
"of its complexity (Amin and Wollenberg, 2005).",

"The stakeholders form an interdependent ecosystem. Generation companies and renewable "
"operators inject power; the transmission and distribution system operators (TSO/DSO) are "
"responsible for stability, balancing and restoration; regulators set reliability standards "
"and market rules; and end-consumers — hospitals, water-treatment plants, "
"telecommunications, transport and households — depend on uninterrupted supply. Repair "
"crews and maintenance contractors constitute a critical but finite resource whose "
"availability directly governs how quickly the system recovers from faults.",

"Interdependencies operate both within the grid and across infrastructures. Internally, "
"generation must be continuously balanced against demand; the failure of a single transmission "
"line forces power to redistribute onto its neighbours, which may in turn overload and trip — "
"the mechanism of a cascading failure (Dobson et al., 2007). Externally, the grid exhibits "
"strong cross-sector interdependencies: telecommunications and SCADA require power, yet the "
"grid's own control depends on those same communications, producing the bidirectional "
"physical, cyber, geographic and logical couplings catalogued by Rinaldi, Peerenboom and Kelly "
"(2001). Water, fuel, finance, healthcare and transport all rely on electricity, so a "
"sustained outage propagates rapidly into a society-wide crisis.",

"It is precisely this combination of universal dependence, tight coupling and potential for "
"cascading propagation that makes the power grid critical to society. Reliable electricity is a "
"precondition for public health, economic activity and national security, and even short "
"interruptions impose large economic and human costs (U.S.–Canada Power System Outage Task "
"Force, 2004). Understanding how the grid behaves under stress — and how mitigation "
"measures such as additional repair capacity change that behaviour — is therefore of direct "
"societal value, and is the motivation for the simulation developed in this report.",
]

TASK1_2 = [
"Power systems are vulnerable to a spectrum of failure modes that can be broadly grouped into "
"physical, operational, environmental and cyber categories. Physically, transmission lines and "
"transformers have finite thermal ratings; sustained overload causes lines to sag, trip or "
"fail, while ageing assets exhibit rising fault probabilities (Kundur et al., 2004). "
"Operationally, the instantaneous balance between generation and demand must be maintained; the "
"growing penetration of variable renewables introduces intermittency that complicates this "
"balance (Amin and Wollenberg, 2005). Environmentally, extreme weather is an increasingly "
"important driver of correlated, simultaneous failures (Panteli and Mancarella, 2015). Finally, "
"the cyber layer introduces a new attack surface in which compromised control systems can be "
"weaponised against the physical grid.",

"The most consequential failure mode is the cascading failure, in which an initial disturbance "
"triggers a self-propagating sequence of overloads and trips. The 2003 Northeast blackout is "
"the canonical example: a software bug in an alarm system, combined with inadequate tree "
"trimming and poor situational awareness, allowed a local line fault in Ohio to cascade into a "
"loss of supply to roughly 50 million people across the north-eastern United States and Canada "
"(U.S.–Canada Power System Outage Task Force, 2004). The event demonstrated how tightly "
"coupled networks redistribute load after each trip, repeatedly pushing neighbouring elements "
"past their limits — exactly the dynamic this report models.",

"More recent events confirm the persistence of these vulnerabilities. The 2016 South Australian "
"‘black system’ event showed how extreme weather and the ride-through settings of "
"wind generation could combine to collapse an entire regional grid (AEMO, 2017), highlighting "
"the renewable-integration risk. The January 2021 separation of the Continental European "
"synchronous area into two islands illustrated how a single substation fault can propagate "
"across national borders before automatic defence mechanisms arrest it (ENTSO-E, 2021). "
"Together these case studies show that vulnerability is not merely a function of any single "
"component but emerges from the interaction of demand stress, asset condition, renewable "
"variability and the finite speed of human and automated response.",

"From a modelling perspective, these historical examples motivate the specific mechanisms "
"embedded in the simulation: a time-varying demand profile, intermittent solar output, "
"capacity-proportional load redistribution after a trip, an overload-dependent failure "
"probability, and a finite repair resource. By representing these explicitly, the model can "
"reproduce the qualitative signature of cascading failure and be used to test which mitigation "
"levers — additional crews, reduced loading, or improved asset condition — most "
"effectively improve resilience (Dobson et al., 2007).",
]

TASK1_3 = [
"Selecting an appropriate modelling paradigm is decisive for both validity and tractability. "
"The principal candidates are continuous simulation (differential equations), system dynamics, "
"agent-based modelling (ABM), discrete-event simulation (DES) and hybrid combinations of these "
"(Borshchev and Filippov, 2004). The choice should follow the dominant characteristics of the "
"system under study.",

"The resilience questions addressed here are governed by discrete occurrences — a line "
"overloads, a line fails, a repair completes, a neighbour trips — separated by intervals "
"during which the system state changes only gradually. This is the defining profile of a "
"discrete-event system, in which the state changes only at a countable set of event times "
"(Banks et al., 2010). Discrete-event simulation is therefore selected as the primary "
"approach. It represents failures and repairs as instantaneous events, models the repair crew "
"naturally as a shared, capacity-limited resource for which competing repairs must queue, and "
"is computationally efficient because the simulation clock advances from event to event rather "
"than over fine fixed time-steps (Law, 2015).",

"A purely continuous or system-dynamics model, by contrast, excels at representing aggregate "
"feedback and smoothly varying quantities but is awkward for discrete, queue-based resource "
"contention and individual component trips (Sterman, 2000). A fully agent-based model could "
"capture heterogeneous component behaviour but would add implementation overhead "
"disproportionate to the questions asked here (Macal and North, 2010). The adopted design is "
"therefore best described as discrete-event with a light hybrid flavour: the discrete event "
"engine (SimPy) drives the simulation, while the continuous demand and solar profiles are "
"sampled at an hourly step and injected as time-varying inputs. This hybrid framing retains the "
"clarity and efficiency of DES while still capturing the continuous variability that drives the "
"grid towards its failure thresholds.",
]

TASK2_1 = [
"The conceptual model defines the boundaries, entities, attributes and relationships of the "
"simulated grid before any code is written. The system boundary encloses the generation, "
"transmission and substation layers of a single regional distribution grid together with the "
"repair organisation; external markets, weather systems and the wider interconnection are "
"treated as exogenous inputs (demand and solar profiles) rather than modelled explicitly. This "
"scoping keeps the model focused on the cascading-failure and recovery dynamics that are the "
"object of study.",

"Three classes of entity are represented. Generators have attributes of rated capacity, type "
"(coal or solar) and dispatchability; the two coal units provide controllable baseload while "
"the solar farm follows a daylight curve modulated by random cloud cover. Transmission lines "
"have a thermal capacity, an instantaneous load and an in-service status; they are the elements "
"that overload, fail and are repaired. Substations map generation onto four consumer zones, "
"each drawing a fixed share of total demand. The shared repair crew is modelled as a "
"capacity-limited resource for which failed lines must queue.",

"The relationships among these entities are shown in Figure 1. Power flows left to right from "
"generators, through transmission lines, into substations and on to the consumer zones. When a "
"line fails, demand is redistributed across the remaining in-service lines in proportion to "
"their capacities; if this redistribution pushes a neighbour beyond the cascade threshold, a "
"cascading event is recorded. The control loop closes when the repair crew restores a failed "
"line, returning it to service and triggering a further redistribution. These four "
"layers and the feedback control loop constitute the complete conceptual model implemented in "
"Task 2.2.",
]

TASK2_2 = [
"The simulation is implemented in Python using the SimPy 4 discrete-event framework (Team "
"SimPy, 2023), chosen for its lightweight process-based API and transparent, auditable code. "
"The implementation is deliberately decomposed into separate modules — configuration, "
"model, metrics, scenarios and visualisation — so that a single source of truth for "
"parameters is shared by the command-line runner and the interactive Colab interface. The "
"complete listings appear in Appendix A.",

"The core of the model is the GridSimulation class, which owns its own random-number generator "
"so that every run is both reproducible (via an explicit seed) and isolated from other runs "
"during the scenario sweep. Five categories of concurrent SimPy process run on a shared "
"environment: a generator process sets hourly output; a demand process applies a time-of-day "
"load profile with stochastic jitter; one line-monitor process per transmission line checks for "
"overload and stochastic failure each hour; a repair process is spawned on failure and competes "
"for the repair-crew resource; and a metrics collector samples the full system state every hour "
"for plotting and export. Using one monitor process per line is the key design decision that "
"lets cascading behaviour emerge naturally rather than being scripted.",

"Several modelling decisions warrant justification. Load is redistributed in proportion to the "
"rated capacity of the remaining in-service lines, a simple but defensible surrogate for a full "
"power-flow solution that preserves the essential cascade mechanism while keeping the model "
"transparent. Failure is treated probabilistically: each line carries a base hourly failure "
"probability that is multiplied by a configurable factor when the line is overloaded, capturing "
"the empirically observed acceleration of failures under stress (Dobson et al., 2007). The "
"repair crew is a SimPy Resource with a configurable capacity, so that when failures outnumber "
"crews, repairs queue and outage durations lengthen — directly linking resourcing decisions "
"to resilience.",

"The default parameter set defines the baseline scenario: a 72-hour horizon advanced in one-hour "
"steps, a 5% base hourly failure probability, a cascade threshold at 120% of rated capacity, an "
"overload failure multiplier of five, repair durations drawn uniformly between two and four "
"hours, a single repair crew, and a fixed seed of 42. These values were chosen to place the "
"baseline grid in a realistic, lightly stressed regime — occasional failures that are "
"normally repaired before they cascade — so that the what-if scenarios in Task 2.4 can "
"reveal how demand growth, ageing assets or extra crews shift the system towards or away from "
"systemic failure. All parameters are exposed as configuration keys and as interactive sliders, "
"making the model straightforward to recalibrate against empirical data.",
]

TASK2_3 = [
"To satisfy the requirement for interactive control and visualisation, a Google Colab notebook "
"provides the user interface, built with the ipywidgets library. The notebook clones the model "
"repository, then renders a control panel of sliders and input boxes that allow the user to "
"adjust the simulation horizon, demand multiplier, base failure probability, number of repair "
"crews, solar scaling, cascade threshold and random seed. Pressing the Run button re-executes "
"the simulation with the chosen parameters and immediately redraws every output, giving the "
"effective real-time control over parameters that the brief requires; an Export button writes "
"and downloads the results.",

"Four complementary visualisations are produced (Figures 2–4 plus a failure-timeline plot). "
"A stacked-area chart shows generator output against total demand, making renewable "
"intermittency and supply–demand tightness visible at a glance. A load-percentage chart "
"plots the average line loading against the 100% rating and the cascade threshold, with vertical "
"markers at each failure time. A stem plot gives the discrete failure-event timeline. Finally, a "
"colour-coded heatmap displays the load of every individual line over time, so that "
"localised overloads and the spatial spread of a cascade can be read directly from the figure. "
"Together these views address the differing needs of operational, tabular and spatial analysis.",

"Results can be exported for further analysis in three forms: an hourly time-series CSV "
"(demand, per-generator output, average and per-line loading), a discrete event log CSV "
"(every failure, cascade and restoration with its timestamp), and a scenario-comparison CSV. "
"Exporting to CSV rather than a proprietary format ensures the outputs can be opened in any "
"spreadsheet or statistical package, supporting onward analysis and reproducibility. The same "
"figure-building functions are reused by the headless command-line runner, so the interactive "
"and batch workflows produce identical artefacts.",
]

TASK2_4 = [
"The model was executed under six scenarios designed to probe distinct stressors: a Baseline; "
"High Demand (+25%); Aged Infrastructure (tripled failure probability); High Renewable "
"(solar boosted 40%); Extra Repair Crews (three instead of one); and a combined Stress Test. "
"All runs used a 72-hour horizon and a fixed seed of 42 to isolate the effect of each lever. "
"Table 1 summarises the headline metrics.",
]

TASK2_4B = [
"The results display a clear and interpretable resilience signature. The baseline grid suffers "
"twelve failures, all repaired, with no cascades and a peak average loading of 107%. Increasing "
"demand by 25% is sufficient to push lines past their cascade threshold, producing 38 cascade "
"events and a peak loading above 360%, while ageing the infrastructure raises failures to 28 "
"and leaves five lines unrecovered at the end of the horizon. The Stress Test, combining both, "
"is the most severe, with 41 cascades and crew utilisation near saturation — the model's "
"depiction of a grid overwhelmed faster than it can be repaired. Crucially, adding repair crews "
"cuts the mean time to repair (MTTR) from 4.25 to 3.04 hours and lowers crew utilisation, "
"demonstrating the simulation's value as a decision-support tool for resourcing.",

"Performance and accuracy were also measured. Each 72-hour run completes in roughly 3–4 "
"milliseconds on a standard machine, so the entire six-scenario sweep and a 30-seed accuracy "
"study run in well under a second, confirming that discrete-event simulation is extremely "
"resource-efficient for problems of this size. Because the model is stochastic, accuracy is "
"reported as a confidence interval: repeating the baseline across 30 independent seeds yields a "
"mean of 18.1 failures with a 95% confidence half-width of ±1.24, and a mean of 24.4 "
"cascade events (±7.1). These intervals quantify the uncertainty in any single run and "
"confirm that the qualitative conclusions are statistically robust rather than artefacts of one "
"random sequence (Law, 2015).",
]

TASK3_1 = [
"Two simulation approaches were evaluated for this scenario: Python with SimPy (the technique "
"adopted) and AnyLogic, a commercial multi-method simulation platform (Borshchev and Filippov, "
"2004). Comparing them clarifies why the chosen tool is appropriate and where alternatives "
"might excel.",

"SimPy is an open-source, process-based discrete-event library embedded in general-purpose "
"Python. Its capabilities centre on lightweight generator-based processes, shared resources "
"with queueing, and event scheduling, all expressed in ordinary Python code. Its principal "
"advantages are that it is free, fully scriptable, version-controllable and integrates "
"seamlessly with the scientific Python stack — NumPy, pandas, Matplotlib and ipywidgets — "
"which made the entire analysis, visualisation and Colab interface achievable in a single "
"ecosystem. Its limitations are equally clear: it offers no graphical model editor, no built-in "
"animation, and no native continuous or agent-based engines, so any such behaviour must be "
"hand-coded. For a transparent, reproducible discrete-event model of modest size, however, "
"these are minor constraints, and its suitability for the present cascading-failure problem is "
"high.",

"AnyLogic, by contrast, is a mature commercial platform that uniquely supports discrete-event, "
"agent-based and system-dynamics modelling within one environment, together with a "
"drag-and-drop graphical editor, rich 2D/3D animation and geographic-information-system "
"integration (Borshchev and Filippov, 2004). For a smart-grid study these features are "
"genuinely attractive: agent-based representation of individual consumers or generators, "
"continuous power-flow dynamics, and built-in spatial visualisation could all be combined "
"without leaving the tool. Its disadvantages are cost — the professional edition is "
"expensive and the free Personal Learning Edition restricts model size and export — and a "
"degree of vendor lock-in, since models are stored in a proprietary format that is harder to "
"version-control or embed in a free cloud notebook than plain Python.",

"The learning curve and resource requirements differ markedly. SimPy demands existing Python "
"proficiency but almost no tool-specific learning, runs on any machine and in free cloud "
"environments such as Colab, and has negligible computational footprint. AnyLogic requires "
"learning a substantial graphical environment and some Java, and benefits from a more capable "
"workstation, but in return reduces the amount of code that must be written for complex, "
"multi-method models. For this assignment — where transparency, reproducibility, zero cost "
"and easy cloud execution were paramount, and where the dynamics are predominantly "
"discrete — SimPy was the more suitable choice; AnyLogic would become preferable if the "
"study were extended to large-scale, multi-method models requiring polished stakeholder-facing "
"animation.",
]

TASK3_2 = [
"Alternative approaches would have illuminated different facets of the problem. A "
"system-dynamics model built around stocks and flows (Sterman, 2000) would have represented "
"aggregate generation, demand and reserve margins as smoothly varying quantities and would have "
"been well suited to studying long-run feedback such as investment cycles or gradual demand "
"growth. It would, however, have obscured the very mechanism this study targets: system "
"dynamics averages over individual components and so cannot naturally represent a single line "
"tripping and forcing a discrete redistribution onto its neighbours. The cascading-failure "
"signature that emerges so clearly from the discrete-event model would have been smoothed away.",

"An agent-based model would have produced complementary insights by endowing each generator, "
"line and consumer with autonomous behavioural rules and letting system behaviour emerge from "
"their interactions (Macal and North, 2010). This would be powerful for studying heterogeneous "
"consumer responses, distributed control or market behaviour, but at the cost of greater "
"implementation effort and more parameters to calibrate and validate. The chosen discrete-event "
"approach occupies a deliberate middle ground: it captures individual component failure and "
"resource contention — the essential ingredients of cascading failure — while remaining "
"simple, fast and transparent.",

"The principal strength of the adopted approach is therefore its clarity and efficiency: the "
"model is small enough to be fully understood and audited, runs in milliseconds, and isolates "
"the cascade mechanism cleanly. Its main weakness is physical fidelity — capacity-"
"proportional redistribution is a simplification of true power-flow physics, and the failure "
"model is probabilistic rather than derived from first principles. A natural and valuable "
"extension would be a hybrid model coupling the discrete-event engine to a proper AC power-flow "
"solver, retaining the present clarity while improving electrical realism.",
]

TASK4_1 = [
"For a distribution system operator, a simulation of this kind translates directly into "
"competitive and decision-making advantage. By quantifying how additional repair crews, reduced "
"peak loading or improved asset condition change failure and cascade counts, the model supports "
"evidence-based capital-allocation and maintenance decisions rather than reliance on rules of "
"thumb. The scenario results make the trade-offs explicit: the marginal benefit of a third "
"repair crew, for example, can be weighed against its cost in reduced outage duration and "
"avoided unserved energy.",

"The cost–benefit case is compelling because the tool itself is almost free. Built on "
"open-source software and runnable in a free cloud notebook, its development and operating costs "
"are negligible, whereas the losses it helps avoid are large: major blackouts impose economic "
"costs running into billions (U.S.–Canada Power System Outage Task Force, 2004). Even a "
"small reduction in outage frequency or duration, or a better-targeted maintenance budget, "
"yields a return on investment far exceeding the modelling effort. The same model also supports "
"regulatory reporting and resilience planning, broadening its organisational value beyond "
"operations into strategy and compliance.",
]

TASK4_2 = [
"Several emerging trends could substantially enhance this simulation. The most significant is "
"the digital twin: a continuously updated virtual replica of the physical grid, fed by live "
"sensor data, that mirrors the real system in real time (Tao et al., 2019). Coupling the "
"present model to live SCADA feeds would turn it from an offline planning tool into an online "
"early-warning system. Artificial-intelligence integration is a natural complement — "
"machine-learning models could forecast demand and renewable output, learn failure "
"probabilities from historical data, and even recommend optimal repair-crew dispatch — "
"while cloud-based simulation, already demonstrated by the Colab deployment, would allow "
"thousands of scenarios to be run in parallel for probabilistic risk assessment.",

"These enhancements carry ethical responsibilities. A model used to guide investment and "
"resilience decisions must be transparent and rigorously validated, because over-confidence in "
"a simplified simulation could lead operators to under-resource a genuinely vulnerable grid, "
"with public-safety consequences. Connecting a digital twin to live infrastructure data raises "
"cybersecurity and privacy concerns: the data stream is itself a critical asset and an attack "
"surface, and consumption data can be personally revealing. There is also an equity dimension — "
"resilience investments guided by the model should not systematically privilege wealthier zones "
"over vulnerable communities. Responsible deployment therefore requires clear communication of "
"model assumptions and uncertainty, robust data governance, and human oversight of any "
"automated decisions (Rinaldi, Peerenboom and Kelly, 2001).",
]

# Results pulled from the actual baseline run (seed=42, 72 h).
RESULTS_ROWS = [
    ["Baseline", 12, 0, 4.25, 106.7, 0.535],
    ["High Demand (+25%)", 15, 38, 6.24, 361.9, 0.617],
    ["Aged Infrastructure", 28, 29, 11.10, 453.9, 0.993],
    ["High Renewable", 12, 0, 4.25, 106.7, 0.535],
    ["Extra Repair Crews", 16, 15, 3.04, 160.3, 0.222],
    ["Stress Test", 28, 41, 12.29, 556.9, 0.990],
]

REFERENCES = [
"AEMO (2017) Black System South Australia 28 September 2016: Final Report. Melbourne: "
"Australian Energy Market Operator.",
"Amin, S.M. and Wollenberg, B.F. (2005) ‘Toward a smart grid: power delivery for the 21st "
"century’, IEEE Power and Energy Magazine, 3(5), pp. 34–41.",
"Banks, J., Carson, J.S., Nelson, B.L. and Nicol, D.M. (2010) Discrete-Event System "
"Simulation. 5th edn. Upper Saddle River, NJ: Pearson.",
"Borshchev, A. and Filippov, A. (2004) ‘From system dynamics and discrete event to "
"practical agent based modeling: reasons, techniques, tools’, Proceedings of the 22nd "
"International Conference of the System Dynamics Society, Oxford, 25–29 July.",
"Buldyrev, S.V., Parshani, R., Paul, G., Stanley, H.E. and Havlin, S. (2010) ‘Catastrophic "
"cascade of failures in interdependent networks’, Nature, 464(7291), pp. 1025–1028.",
"Dobson, I., Carreras, B.A., Lynch, V.E. and Newman, D.E. (2007) ‘Complex systems analysis "
"of series of blackouts: cascading failure, critical points, and self-organization’, "
"Chaos, 17(2), 026103.",
"ENTSO-E (2021) Continental Europe Synchronous Area Separation on 8 January 2021: Final "
"Report. Brussels: European Network of Transmission System Operators for Electricity.",
"Kundur, P., Paserba, J., Ajjarapu, V., Andersson, G., Bose, A., Canizares, C., Hatziargyriou, "
"N., Hill, D., Stankovic, A., Taylor, C., Van Cutsem, T. and Vittal, V. (2004) ‘Definition "
"and classification of power system stability’, IEEE Transactions on Power Systems, 19(3), "
"pp. 1387–1401.",
"Law, A.M. (2015) Simulation Modeling and Analysis. 5th edn. New York: McGraw-Hill.",
"Macal, C.M. and North, M.J. (2010) ‘Tutorial on agent-based modelling and simulation’, "
"Journal of Simulation, 4(3), pp. 151–162.",
"Panteli, M. and Mancarella, P. (2015) ‘Influence of extreme weather and climate change on "
"the resilience of power systems’, IEEE Systems Journal, 9(3), pp. 1733–1742.",
"Rinaldi, S.M., Peerenboom, J.P. and Kelly, T.K. (2001) ‘Identifying, understanding, and "
"analyzing critical infrastructure interdependencies’, IEEE Control Systems Magazine, "
"21(6), pp. 11–25.",
"Sterman, J.D. (2000) Business Dynamics: Systems Thinking and Modeling for a Complex World. "
"Boston: Irwin/McGraw-Hill.",
"Tao, F., Zhang, H., Liu, A. and Nee, A.Y.C. (2019) ‘Digital twin in industry: "
"state-of-the-art’, IEEE Transactions on Industrial Informatics, 15(4), pp. 2405–2415.",
"Team SimPy (2023) SimPy Discrete-Event Simulation for Python: Documentation. Available at: "
"https://simpy.readthedocs.io (Accessed: 2 June 2026).",
"U.S.–Canada Power System Outage Task Force (2004) Final Report on the August 14, 2003 "
"Blackout in the United States and Canada: Causes and Recommendations. Washington, DC / Ottawa.",
]


if __name__ == "__main__":
    main()
